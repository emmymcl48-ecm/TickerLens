"""Generate TickerLens technical documentation as a .docx file."""
import docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# -- Styles --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

# ── Title ──
title = doc.add_heading('TickerLens — Technical Documentation', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('AI-Powered Equity Analysis Platform\nFIN 372 — AI and Portfolio Management\nThe University of Texas at Austin')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(100, 100, 100)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub2.add_run('Anissa de la Mora, Emily McLaughlin, Elena Cokinos, & Abbie Kate Henderson')
run2.font.size = Pt(11)
run2.font.italic = True

doc.add_page_break()

# ── Table of Contents ──
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Project Overview',
    '2. Architecture & Technology Stack',
    '3. Data Sources & Retrieval Strategy',
    '   3.1 Live Financial Metrics (yfinance + Claude Fallback)',
    '   3.2 Fama-French 48 Industry Medians',
    '   3.3 Earnings Call Analysis',
    '   3.4 Data Source Attribution',
    '4. Scoring Methodology',
    '   4.1 Valuation Pillar',
    '   4.2 Profitability Pillar',
    '   4.3 Risk & Momentum Pillar',
    '   4.4 Earnings Sentiment Pillar',
    '   4.5 Overall Signal Computation',
    '5. Sector-Relative Benchmarking (Fama-French 48)',
    '6. AI Integration — Claude API',
    '7. Frontend Design (Streamlit)',
    '8. Deployment & Infrastructure',
    '9. Rate Limiting & Resilience',
    '10. Parameters & Configuration Reference',
    '11. Example Outputs',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ═══════════════════════════════════════════
# 1. PROJECT OVERVIEW
# ═══════════════════════════════════════════
doc.add_heading('1. Project Overview', level=1)

doc.add_paragraph(
    'TickerLens is a web-based equity analysis platform that provides instant, '
    'structured investment signals for any publicly traded stock. A user enters '
    'a stock ticker and receives a comprehensive, non-biased assessment across '
    'four analytical pillars: Valuation, Profitability, Risk & Momentum, and '
    'Earnings Sentiment.'
)

doc.add_paragraph(
    'Each pillar produces a score from 1 to 10 and a signal — Green (bullish), '
    'Yellow (neutral), or Red (bearish). These pillar scores are combined using '
    'a weighted average to produce an overall investment verdict. The tool is '
    'designed to be objective, comparing each stock against its own industry '
    'sector rather than using universal thresholds.'
)

doc.add_heading('Key Design Principles', level=2)
bullets = [
    'Non-biased: Scores are computed algorithmically against sector medians, not subjective opinion.',
    'Sector-relative: A P/E of 30 means something different in tech vs. utilities. TickerLens accounts for this.',
    'Transparent: Every score shows the underlying data comparison, and all data sources are cited with retrieval dates.',
    'Real-time: All financial data is fetched live at the time of analysis using today\'s market data.',
    'AI-enhanced: Earnings call transcripts are found via web search and analyzed by Claude AI for sentiment and forward guidance.',
    'Resilient: Multiple data retrieval strategies ensure the tool works reliably across local and cloud environments.',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# ═══════════════════════════════════════════
# 2. ARCHITECTURE
# ═══════════════════════════════════════════
doc.add_heading('2. Architecture & Technology Stack', level=1)

doc.add_paragraph(
    'TickerLens is built as a Python Streamlit application. Streamlit provides '
    'both the web server and the interactive frontend in a single framework, '
    'simplifying deployment and maintenance.'
)

doc.add_heading('Technology Stack', level=2)
table = doc.add_table(rows=9, cols=2)
table.style = 'Light Shading Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = table.rows[0].cells
headers[0].text = 'Component'
headers[1].text = 'Technology'
data = [
    ('Language', 'Python 3.x'),
    ('Web Framework', 'Streamlit (frontend + server in one)'),
    ('Financial Data (Primary)', 'yfinance — Python Yahoo Finance library with curl_cffi'),
    ('Financial Data (Fallback)', 'Claude API with web search (when yfinance is blocked on cloud IPs)'),
    ('Earnings Analysis', 'Claude API (Anthropic) with web search — claude-haiku-4-5 model'),
    ('Sector Benchmarks', 'Pre-computed JSON from WRDS/Compustat (Fama-French 48 industries)'),
    ('Industry Mapping', 'Keyword-based Yahoo Finance industry to FF48 lookup'),
    ('Deployment', 'Render.com (free tier, auto-deploy from GitHub)'),
]
for i, (comp, tech) in enumerate(data):
    row = table.rows[i + 1].cells
    row[0].text = comp
    row[1].text = tech

doc.add_heading('Project File Structure', level=2)
files = [
    ('app.py', 'Main Streamlit application — UI, pipeline orchestration, result rendering'),
    ('modules_py/fetch_metrics.py', 'Fetches live financial metrics via yfinance (primary)'),
    ('modules_py/claude_combined.py', 'Combined Claude fallback — fetches metrics + sentiment in one API call'),
    ('modules_py/earnings_sentiment.py', 'Standalone earnings sentiment via Claude web search (used when yfinance works)'),
    ('modules_py/scoring.py', 'Quantitative scoring engine — compares metrics to sector medians'),
    ('modules_py/industry_map.py', 'Maps Yahoo Finance industry names to Fama-French 48 sectors'),
    ('data/sectorData.json', 'Pre-computed FF48 industry median financial ratios (from WRDS)'),
    ('.streamlit/config.toml', 'Streamlit theme configuration (dark finance theme)'),
    ('requirements.txt', 'Python dependencies'),
    ('.env', 'Environment variables — ANTHROPIC_API_KEY (not committed to Git)'),
]
table2 = doc.add_table(rows=len(files) + 1, cols=2)
table2.style = 'Light Shading Accent 1'
table2.rows[0].cells[0].text = 'File'
table2.rows[0].cells[1].text = 'Purpose'
for i, (f, desc) in enumerate(files):
    table2.rows[i + 1].cells[0].text = f
    table2.rows[i + 1].cells[1].text = desc

# ═══════════════════════════════════════════
# 3. DATA SOURCES
# ═══════════════════════════════════════════
doc.add_heading('3. Data Sources & Retrieval Strategy', level=1)

doc.add_paragraph(
    'TickerLens uses a multi-layered data retrieval strategy designed to always '
    'return current, real-time data regardless of the deployment environment.'
)

doc.add_heading('3.1 Live Financial Metrics (yfinance + Claude Fallback)', level=2)

doc.add_paragraph(
    'Financial metrics are always retrieved live — never cached or stale. The '
    'system uses a two-tier approach:'
)

p = doc.add_paragraph()
run = p.add_run('Primary — yfinance: ')
run.bold = True
p.add_run(
    'The Python yfinance library (with curl_cffi for cloud compatibility) '
    'fetches data directly from Yahoo Finance. This is fast, free, and provides '
    'comprehensive financial data. Works reliably on local machines and many '
    'cloud environments.'
)

p = doc.add_paragraph()
run = p.add_run('Fallback — Claude Web Search: ')
run.bold = True
p.add_run(
    'When yfinance is blocked (Yahoo rate-limits some cloud server IPs), the '
    'system falls back to a combined Claude API call with web search enabled. '
    'Claude searches financial websites (Yahoo Finance, Google Finance, '
    'MarketWatch, etc.) for current data and returns structured metrics. This '
    'fallback is combined with the earnings sentiment analysis in a single API '
    'call to avoid rate limiting.'
)

doc.add_paragraph('The following metrics are retrieved for each ticker:')

metrics_list = [
    ('Current Price', 'Latest market price'),
    ('Market Cap', 'Total market capitalization'),
    ('Trailing P/E', 'Price / trailing 12-month earnings per share'),
    ('Forward P/E', 'Price / estimated next-12-month EPS'),
    ('Price-to-Book', 'Market price / book value per share'),
    ('EV/EBITDA', 'Enterprise value / EBITDA'),
    ('EV/Revenue', 'Enterprise value / total revenue'),
    ('Gross Margin', 'Gross profit / revenue (decimal)'),
    ('Operating Margin', 'Operating income / revenue (decimal)'),
    ('Net Margin', 'Net income / revenue (decimal)'),
    ('Return on Equity', 'Net income / shareholders\' equity (decimal)'),
    ('Return on Assets', 'Net income / total assets (decimal)'),
    ('Current Ratio', 'Current assets / current liabilities'),
    ('Quick Ratio', '(Current assets - inventory) / current liabilities'),
    ('Debt-to-Equity', 'Total debt / shareholders\' equity (percentage)'),
    ('Beta', 'Systematic risk relative to the market'),
    ('50-Day Avg', '50-day simple moving average of price'),
    ('200-Day Avg', '200-day simple moving average of price'),
    ('52-Week High/Low', 'Highest and lowest prices over past 52 weeks'),
    ('Revenue Growth', 'Year-over-year revenue growth rate (decimal)'),
    ('Earnings Growth', 'Year-over-year earnings growth rate (decimal)'),
]
table3 = doc.add_table(rows=len(metrics_list) + 1, cols=2)
table3.style = 'Light Shading Accent 1'
table3.rows[0].cells[0].text = 'Metric'
table3.rows[0].cells[1].text = 'Definition'
for i, (m, d) in enumerate(metrics_list):
    table3.rows[i + 1].cells[0].text = m
    table3.rows[i + 1].cells[1].text = d

doc.add_heading('3.2 Fama-French 48 Industry Medians', level=2)
doc.add_paragraph(
    'To enable sector-relative scoring, we use pre-computed median financial '
    'ratios for each of the Fama-French 48 industry classifications. This data '
    'was sourced from WRDS (Wharton Research Data Services) and covers the period '
    '2015-2025. For each industry, we use the most recent year\'s medians.'
)

doc.add_paragraph('The sector median metrics available are:')
sector_metrics = [
    ('bm', 'Book-to-Market ratio (median)'),
    ('evm', 'Enterprise Value Multiple / EBITDA (median)'),
    ('pe_inc', 'Price-to-Earnings including extraordinary items (median)'),
    ('pe_op_dil', 'Price-to-Earnings from operations, diluted (median)'),
    ('GProf', 'Gross Profitability = gross profit / assets (median)'),
    ('aftret_invcapx', 'After-tax return on invested capital (median)'),
    ('gpm', 'Gross Profit Margin (median)'),
    ('npm', 'Net Profit Margin (median)'),
    ('quick_ratio', 'Quick Ratio (median)'),
]
table4 = doc.add_table(rows=len(sector_metrics) + 1, cols=2)
table4.style = 'Light Shading Accent 1'
table4.rows[0].cells[0].text = 'Field'
table4.rows[0].cells[1].text = 'Description'
for i, (f, d) in enumerate(sector_metrics):
    table4.rows[i + 1].cells[0].text = f
    table4.rows[i + 1].cells[1].text = d

doc.add_heading('3.3 Earnings Call Analysis', level=2)
doc.add_paragraph(
    'TickerLens uses Claude AI with web search to find and analyze the most '
    'recent quarterly earnings call for each ticker. This is done in a single '
    'API call that both searches for the earnings information and produces the '
    'sentiment analysis. The search prioritizes company investor relations pages '
    'first, then financial news sources like Motley Fool and Seeking Alpha.'
)

doc.add_paragraph(
    'The earnings sentiment output includes the fiscal quarter and date of the '
    'earnings call (e.g., "[Q4 2025 - January 28, 2026]"), so users know exactly '
    'which earnings report is being analyzed.'
)

doc.add_heading('3.4 Data Source Attribution', level=2)
doc.add_paragraph(
    'Every analysis displays a "Data Sources" section showing:'
)
sources = [
    'Financial Metrics: The source (Yahoo Finance or Claude Web Search) and the date retrieved.',
    'Sector Benchmarks: WRDS / Compustat (Fama-French 48 Industry Medians).',
    'Earnings Sentiment: Claude AI (Anthropic) with web search, including the earnings quarter and call date.',
]
for s in sources:
    doc.add_paragraph(s, style='List Bullet')

# ═══════════════════════════════════════════
# 4. SCORING METHODOLOGY
# ═══════════════════════════════════════════
doc.add_heading('4. Scoring Methodology', level=1)

doc.add_paragraph(
    'Each pillar scores the stock on a 1-10 scale. Individual metric scores are '
    'computed by comparing the stock\'s value to its sector median, then averaged '
    'within each pillar. The core scoring formula is:'
)

doc.add_paragraph(
    'For "lower is better" metrics (e.g., P/E, EV/EBITDA):\n'
    '    score = 10 - (stock_value / sector_median) x 5\n\n'
    'For "higher is better" metrics (e.g., margins, ROE):\n'
    '    score = (stock_value / sector_median) x 5\n\n'
    'All scores are clamped to the range [1, 10].'
)

doc.add_paragraph(
    'Signal thresholds:\n'
    '    Green (Bullish):  score >= 6.5\n'
    '    Yellow (Neutral): score >= 4.0 and < 6.5\n'
    '    Red (Bearish):    score < 4.0'
)

doc.add_heading('4.1 Valuation Pillar', level=2)
doc.add_paragraph(
    'Measures whether the stock is cheap or expensive relative to its sector. '
    'Lower valuations score higher.'
)
val_metrics = [
    ('Trailing P/E vs. sector median P/E', 'Lower is better', 'Equal weight'),
    ('Book-to-Market vs. sector median B/M', 'Higher is better (cheaper)', 'Equal weight'),
    ('EV/EBITDA vs. sector median EVM', 'Lower is better', 'Equal weight'),
]
table5 = doc.add_table(rows=len(val_metrics) + 1, cols=3)
table5.style = 'Light Shading Accent 1'
table5.rows[0].cells[0].text = 'Comparison'
table5.rows[0].cells[1].text = 'Direction'
table5.rows[0].cells[2].text = 'Weight'
for i, row_data in enumerate(val_metrics):
    for j, val in enumerate(row_data):
        table5.rows[i + 1].cells[j].text = val

doc.add_heading('4.2 Profitability Pillar', level=2)
doc.add_paragraph(
    'Measures the company\'s ability to generate profits relative to its sector peers. '
    'Higher profitability scores higher.'
)
prof_metrics = [
    ('Gross Margin vs. sector median', 'Higher is better', 'Equal weight'),
    ('Net Margin vs. sector median', 'Higher is better', 'Equal weight'),
    ('Return on Equity (absolute)', '> 20% is strong', 'Equal weight'),
    ('Revenue Growth (absolute)', '> 10% is strong', 'Equal weight'),
]
table6 = doc.add_table(rows=len(prof_metrics) + 1, cols=3)
table6.style = 'Light Shading Accent 1'
table6.rows[0].cells[0].text = 'Metric'
table6.rows[0].cells[1].text = 'Direction'
table6.rows[0].cells[2].text = 'Weight'
for i, row_data in enumerate(prof_metrics):
    for j, val in enumerate(row_data):
        table6.rows[i + 1].cells[j].text = val

doc.add_paragraph(
    'Note: ROE and Revenue Growth are scored on absolute scales rather than '
    'vs. sector median, since these metrics are universally interpretable.'
)

doc.add_heading('4.3 Risk & Momentum Pillar', level=2)
doc.add_paragraph(
    'Assesses financial health, risk exposure, and price momentum.'
)
risk_metrics = [
    ('Quick Ratio vs. sector median', 'Higher is better (more liquid)', 'Equal weight'),
    ('Debt-to-Equity (absolute)', 'Lower is better: score = 9 - (D/E / 30)', 'Equal weight'),
    ('Beta (absolute)', 'Closer to 1 is neutral: score = 8 - (beta-0.5) x 3', 'Equal weight'),
    ('Price vs. 200-Day Avg', 'Above = positive momentum: score = 5 + (% above) x 15', 'Equal weight'),
]
table7 = doc.add_table(rows=len(risk_metrics) + 1, cols=3)
table7.style = 'Light Shading Accent 1'
table7.rows[0].cells[0].text = 'Metric'
table7.rows[0].cells[1].text = 'Scoring Logic'
table7.rows[0].cells[2].text = 'Weight'
for i, row_data in enumerate(risk_metrics):
    for j, val in enumerate(row_data):
        table7.rows[i + 1].cells[j].text = val

doc.add_heading('4.4 Earnings Sentiment Pillar', level=2)
doc.add_paragraph(
    'This pillar uses Claude AI to search the web for the most recent earnings '
    'call and analyze it in a single step. The search prioritizes company investor '
    'relations pages first, then financial news sources. Claude evaluates:'
)
sentiment_factors = [
    'Revenue and earnings guidance — was it raised, maintained, or lowered?',
    'Management confidence vs. hedging language',
    'Key growth drivers or headwinds mentioned',
    'Surprises or strategic changes',
    'Overall tone — confident, cautious, or defensive',
]
for s in sentiment_factors:
    doc.add_paragraph(s, style='List Bullet')
doc.add_paragraph(
    'Claude returns a score (1-10), a signal (Green/Yellow/Red), the fiscal quarter '
    'and earnings call date, and a 2-3 sentence plain-English reasoning. The prompt '
    'explicitly instructs Claude to be objective and non-biased, focusing on what '
    'management actually said rather than market reaction.'
)

doc.add_heading('4.5 Overall Signal Computation', level=2)
doc.add_paragraph(
    'The four pillar scores are combined using a weighted average:'
)
weights = [
    ('Valuation', '30%'),
    ('Profitability', '30%'),
    ('Risk & Momentum', '20%'),
    ('Earnings Sentiment', '20%'),
]
table8 = doc.add_table(rows=len(weights) + 1, cols=2)
table8.style = 'Light Shading Accent 1'
table8.rows[0].cells[0].text = 'Pillar'
table8.rows[0].cells[1].text = 'Weight'
for i, (p, w) in enumerate(weights):
    table8.rows[i + 1].cells[0].text = p
    table8.rows[i + 1].cells[1].text = w

doc.add_paragraph(
    'Rationale for weights: Valuation and Profitability are weighted most heavily (30% each) '
    'because they represent the fundamental investment case. Risk & Momentum and Earnings '
    'Sentiment each get 20% as they provide important but supplementary context.'
)

# ═══════════════════════════════════════════
# 5. SECTOR-RELATIVE BENCHMARKING
# ═══════════════════════════════════════════
doc.add_heading('5. Sector-Relative Benchmarking (Fama-French 48)', level=1)

doc.add_paragraph(
    'A critical design decision in TickerLens is to score stocks relative to their '
    'own industry sector rather than using universal thresholds. This is important '
    'because financial characteristics vary dramatically across industries:'
)

examples = [
    'Software companies typically have P/E ratios of 30-50x, while banks trade at 10-15x.',
    'Grocery retailers have net margins of 1-3%, while pharmaceutical companies have 15-25%.',
    'Utilities have high D/E ratios by design, while tech companies often carry little debt.',
]
for e in examples:
    doc.add_paragraph(e, style='List Bullet')

doc.add_paragraph(
    'We use the Fama-French 48 industry classification system, which groups all public '
    'companies into 48 industries based on SIC codes. For each industry, we computed '
    'median values of key financial ratios using data from WRDS (2015-2025). When a '
    'stock is analyzed, we identify its FF48 industry via a keyword-based mapping from '
    'the Yahoo Finance industry name (e.g., "Semiconductors" maps to CHIPS, '
    '"Auto Manufacturers" maps to AUTOS) and compare its metrics against '
    'the corresponding medians.'
)

# ═══════════════════════════════════════════
# 6. AI INTEGRATION
# ═══════════════════════════════════════════
doc.add_heading('6. AI Integration — Claude API', level=1)

doc.add_paragraph('Claude AI (by Anthropic) is used in TickerLens in two configurations:')

doc.add_heading('Path A: yfinance Available (Local / Compatible Cloud)', level=2)
doc.add_paragraph(
    'When yfinance successfully retrieves financial metrics, a single Claude API call '
    'with web search is made for earnings sentiment analysis only. Claude searches '
    'for the most recent earnings call, starting with the company\'s investor relations '
    'page, then financial news sources. It returns a JSON object with the sentiment '
    'score, signal, quarter, date, and reasoning.'
)

doc.add_heading('Path B: yfinance Blocked (Cloud Fallback)', level=2)
doc.add_paragraph(
    'When yfinance is blocked by Yahoo (common on cloud server IPs), a single combined '
    'Claude API call with web search retrieves both current financial metrics AND '
    'earnings sentiment analysis in one request. This approach was specifically designed '
    'to avoid rate limiting — making two separate web search calls back-to-back would '
    'exceed the API rate limit of 50,000 input tokens per minute.'
)

doc.add_paragraph(
    'In both paths, only one Claude web search call is ever made per analysis. '
    'The model used is claude-haiku-4-5 (fast, cost-effective). Web search is enabled '
    'via Anthropic\'s built-in web_search tool, limited to 5 searches per call.'
)

# ═══════════════════════════════════════════
# 7. FRONTEND
# ═══════════════════════════════════════════
doc.add_heading('7. Frontend Design (Streamlit)', level=1)

doc.add_paragraph(
    'The frontend is built entirely with Streamlit, using custom CSS for a dark '
    'finance-themed design. Key UI elements:'
)

frontend_items = [
    'Ticker input field with uppercase formatting and Enter-key support',
    'Loading spinner with estimated wait time ("Analyzing... this may take 15-30 seconds")',
    'Overall verdict banner showing company name, price, market cap, sector, and composite signal',
    'Four pillar cards in a 2x2 grid, each displaying: pillar name, numeric score with color, '
    'animated score bar, signal badge (Green/Yellow/Red), and plain-English reasoning',
    'Data Sources section showing where metrics came from and when they were retrieved',
    'Error handling with descriptive messages for invalid tickers or API failures',
]
for item in frontend_items:
    doc.add_paragraph(item, style='List Bullet')

# ═══════════════════════════════════════════
# 8. DEPLOYMENT
# ═══════════════════════════════════════════
doc.add_heading('8. Deployment & Infrastructure', level=1)

doc.add_paragraph(
    'TickerLens is deployed on Render.com as a free-tier web service, connected to '
    'the GitHub repository for automatic deployments on every push to the main branch.'
)

deploy_details = [
    ('GitHub Repository', 'github.com/emmymcl48-ecm/TickerLens'),
    ('Hosting Platform', 'Render.com (free tier)'),
    ('Build Command', 'pip install -r requirements.txt'),
    ('Start Command', 'streamlit run app.py'),
    ('Environment Variables', 'ANTHROPIC_API_KEY (set in Render dashboard, not in code)'),
    ('Auto-Deploy', 'Yes — triggers on push to main branch'),
    ('Python Version', '3.x (determined by Render runtime)'),
]
table9 = doc.add_table(rows=len(deploy_details) + 1, cols=2)
table9.style = 'Light Shading Accent 1'
table9.rows[0].cells[0].text = 'Setting'
table9.rows[0].cells[1].text = 'Value'
for i, (s, v) in enumerate(deploy_details):
    table9.rows[i + 1].cells[0].text = s
    table9.rows[i + 1].cells[1].text = v

# ═══════════════════════════════════════════
# 9. RATE LIMITING & RESILIENCE
# ═══════════════════════════════════════════
doc.add_heading('9. Rate Limiting & Resilience', level=1)

doc.add_paragraph(
    'TickerLens was designed to handle two common failure modes in cloud deployment:'
)

doc.add_heading('Yahoo Finance Blocking', level=2)
doc.add_paragraph(
    'Yahoo Finance aggressively rate-limits and blocks requests from cloud server '
    'IPs (returning HTTP 429 or HTML consent pages instead of JSON). TickerLens '
    'handles this by detecting yfinance failures and automatically falling back '
    'to the combined Claude web search approach. The fallback retrieves current '
    'metrics from alternative financial sources (Google Finance, MarketWatch, etc.) '
    'via Claude\'s web search capability.'
)

doc.add_heading('Anthropic API Rate Limits', level=2)
doc.add_paragraph(
    'The Anthropic API enforces a rate limit of 50,000 input tokens per minute '
    'on standard plans. Web search calls are particularly expensive because search '
    'results are included in the input token count. TickerLens manages this by:'
)
rate_items = [
    'Never making more than one web search call per analysis (combined call when yfinance fails).',
    'Automatic retry with 10-15 second delay when rate limited.',
    'Reduced search budget (2 searches instead of 5) on retry attempts.',
    'Graceful degradation — always returns a result, even if only with a neutral default.',
]
for item in rate_items:
    doc.add_paragraph(item, style='List Bullet')

# ═══════════════════════════════════════════
# 10. PARAMETERS REFERENCE
# ═══════════════════════════════════════════
doc.add_heading('10. Parameters & Configuration Reference', level=1)

doc.add_heading('Scoring Parameters', level=2)
params = [
    ('Green threshold', '>= 6.5', 'Score at or above this = bullish signal'),
    ('Yellow threshold', '>= 4.0', 'Score between 4.0 and 6.5 = neutral signal'),
    ('Red threshold', '< 4.0', 'Score below this = bearish signal'),
    ('Valuation weight', '0.30', 'Weight of Valuation pillar in overall score'),
    ('Profitability weight', '0.30', 'Weight of Profitability pillar in overall score'),
    ('Risk & Momentum weight', '0.20', 'Weight of Risk & Momentum pillar in overall score'),
    ('Earnings Sentiment weight', '0.20', 'Weight of Earnings Sentiment pillar in overall score'),
    ('Score range', '1-10', 'All individual metric scores clamped to this range'),
]
table10 = doc.add_table(rows=len(params) + 1, cols=3)
table10.style = 'Light Shading Accent 1'
table10.rows[0].cells[0].text = 'Parameter'
table10.rows[0].cells[1].text = 'Value'
table10.rows[0].cells[2].text = 'Description'
for i, row_data in enumerate(params):
    for j, val in enumerate(row_data):
        table10.rows[i + 1].cells[j].text = val

doc.add_heading('Metric Scoring Formulas', level=2)
formulas = [
    ('P/E vs sector', 'score = 10 - (PE / median_PE) x 5', 'Lower is better'),
    ('B/M vs sector', 'score = (BM / median_BM) x 5', 'Higher is better (cheaper)'),
    ('EV/EBITDA vs sector', 'score = 10 - (EV / median_EV) x 5', 'Lower is better'),
    ('Gross Margin vs sector', 'score = (GM / median_GM) x 5', 'Higher is better'),
    ('Net Margin vs sector', 'score = (NM / median_NM) x 5', 'Higher is better'),
    ('ROE (absolute)', 'score = ROE x 25 (clamped 1-10)', '4% -> 1, 20% -> 5, 40% -> 10'),
    ('Revenue Growth (absolute)', 'score = 5 + growth x 10', '0% -> 5, 20% -> 7, 50% -> 10'),
    ('Debt/Equity (absolute)', 'score = 9 - (D/E / 30)', '0% -> 9, 100% -> 6, 200% -> 2'),
    ('Beta (absolute)', 'score = 8 - (beta - 0.5) x 3', '0.5 -> 8, 1.0 -> 6.5, 2.0 -> 3.5'),
    ('Price Momentum', 'score = 5 + (% above 200d avg) x 15', '0% -> 5, +20% -> 8, -20% -> 2'),
]
table11 = doc.add_table(rows=len(formulas) + 1, cols=3)
table11.style = 'Light Shading Accent 1'
table11.rows[0].cells[0].text = 'Metric'
table11.rows[0].cells[1].text = 'Formula'
table11.rows[0].cells[2].text = 'Interpretation'
for i, row_data in enumerate(formulas):
    for j, val in enumerate(row_data):
        table11.rows[i + 1].cells[j].text = val

doc.add_heading('Fallback Medians', level=2)
doc.add_paragraph(
    'When a stock\'s industry cannot be mapped to an FF48 sector, the following '
    'fallback medians are used (approximate broad-market values):'
)
fallbacks = [
    ('P/E', '20'),
    ('Book-to-Market', '0.4'),
    ('EV/EBITDA', '12'),
    ('Gross Margin', '30%'),
    ('Net Margin', '8%'),
    ('Quick Ratio', '1.2'),
]
table12 = doc.add_table(rows=len(fallbacks) + 1, cols=2)
table12.style = 'Light Shading Accent 1'
table12.rows[0].cells[0].text = 'Metric'
table12.rows[0].cells[1].text = 'Fallback Median'
for i, (m, v) in enumerate(fallbacks):
    table12.rows[i + 1].cells[0].text = m
    table12.rows[i + 1].cells[1].text = v

# ═══════════════════════════════════════════
# 11. EXAMPLE OUTPUTS
# ═══════════════════════════════════════════
doc.add_heading('11. Example Outputs', level=1)

doc.add_heading('NVIDIA (NVDA) — Strong Stock', level=2)
doc.add_paragraph('Overall: Green (6.6/10)')
nvda_results = [
    ('Valuation', 'Red (2/10)', 'P/E 38.5 vs sector 31.2; EV/EBITDA 34.0 vs sector 12.3 — expensive even for semis'),
    ('Profitability', 'Green (9.5/10)', 'Gross margin 71.1% vs sector 42.6%; net margin 55.6% vs sector 1.6% — dominant'),
    ('Risk & Momentum', 'Green (6.5/10)', 'Quick ratio 3.14 vs sector 1.71; low debt; high beta (2.33)'),
    ('Earnings Sentiment', 'Green (9/10)', '[Q4 2025] Exceptional confidence; 73% YoY revenue growth; strong forward guidance'),
]
table13 = doc.add_table(rows=len(nvda_results) + 1, cols=3)
table13.style = 'Light Shading Accent 1'
table13.rows[0].cells[0].text = 'Pillar'
table13.rows[0].cells[1].text = 'Signal'
table13.rows[0].cells[2].text = 'Key Reasoning'
for i, row_data in enumerate(nvda_results):
    for j, val in enumerate(row_data):
        table13.rows[i + 1].cells[j].text = val

doc.add_heading('Meta Platforms (META) — Mixed Signal', level=2)
doc.add_paragraph('Overall: Yellow (5.7/10)')
meta_results = [
    ('Valuation', 'Red (2.3/10)', 'P/E 27 vs sector; high EV/EBITDA relative to internet peers'),
    ('Profitability', 'Green (8.3/10)', 'Strong margins and ROE; 24% revenue growth'),
    ('Risk & Momentum', 'Yellow (5/10)', 'Moderate debt; high capex guidance ($115-135B) adds risk'),
    ('Earnings Sentiment', 'Green (7.5/10)', '[Q4 2025 - Jan 28, 2026] Beat estimates; raised guidance; cautious on near-term AI monetization'),
]
table14 = doc.add_table(rows=len(meta_results) + 1, cols=3)
table14.style = 'Light Shading Accent 1'
table14.rows[0].cells[0].text = 'Pillar'
table14.rows[0].cells[1].text = 'Signal'
table14.rows[0].cells[2].text = 'Key Reasoning'
for i, row_data in enumerate(meta_results):
    for j, val in enumerate(row_data):
        table14.rows[i + 1].cells[j].text = val

doc.add_heading('Ford (F) — Weak Stock', level=2)
doc.add_paragraph('Overall: Red (3.8/10)')
f_results = [
    ('Valuation', 'Red (3.5/10)', 'Slightly expensive for auto sector'),
    ('Profitability', 'Red (2.3/10)', 'Thin margins, weak ROE vs auto peers'),
    ('Risk & Momentum', 'Red (3.5/10)', 'High debt, negative price momentum'),
    ('Earnings Sentiment', 'Green (6.8/10)', '[Q4 2025] Solid US market share gains, but $4B+ EV losses weigh on outlook'),
]
table15 = doc.add_table(rows=len(f_results) + 1, cols=3)
table15.style = 'Light Shading Accent 1'
table15.rows[0].cells[0].text = 'Pillar'
table15.rows[0].cells[1].text = 'Signal'
table15.rows[0].cells[2].text = 'Key Reasoning'
for i, row_data in enumerate(f_results):
    for j, val in enumerate(row_data):
        table15.rows[i + 1].cells[j].text = val

doc.add_paragraph(
    'These examples demonstrate that TickerLens produces differentiated, plausible '
    'results — correctly identifying NVIDIA\'s profitability dominance offset by '
    'expensive valuation, Meta\'s strong earnings with high-risk capex spending, '
    'and Ford\'s fundamental weakness despite decent earnings sentiment.'
)

# ── Save ──
output_path = r'C:\Users\emmym\TickerLens\TickerLens_Documentation.docx'
doc.save(output_path)
print(f'Document saved to {output_path}')
